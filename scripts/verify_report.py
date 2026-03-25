# -*- coding: utf-8 -*-
"""
Word 报告数据一致性验证脚本
用法：python scripts/verify_report.py "报告.docx" "数据.xlsx"
输出：检查报告中的数字是否与数据一致
"""
import sys
import os
import re
import numpy as np
import pandas as pd

sys.stdout.reconfigure(encoding='utf-8')

def verify_report(docx_path, data_path=None):
    """验证 Word 报告的数据一致性"""
    from docx import Document
    
    print(f'═══ 报告数据一致性验证 ═══\n')
    
    # 1. 提取报告中所有文字
    doc = Document(docx_path)
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    
    # 也提取表格中的文字
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_texts.append(cell.text)
    all_table_text = '\n'.join(table_texts)
    
    print(f'报告: {docx_path}')
    print(f'段落数: {len(doc.paragraphs)}')
    print(f'表格数: {len(doc.tables)}')
    
    # 2. 提取所有数字
    numbers_in_text = re.findall(r'\d+\.?\d*', all_text)
    numbers_in_tables = re.findall(r'\d+\.?\d*', all_table_text)
    
    print(f'文字中的数字: {len(numbers_in_text)} 个')
    print(f'表格中的数字: {len(numbers_in_tables)} 个')
    
    # 3. 检查样本量一致性
    print(f'\n─── 样本量一致性检查 ───')
    n_pattern = re.findall(r'[nN]\s*[=＝]\s*(\d+)', all_text + all_table_text)
    if n_pattern:
        n_values = [int(n) for n in n_pattern]
        unique_ns = set(n_values)
        if len(unique_ns) == 1:
            print(f'  ✅ 样本量一致: N={list(unique_ns)[0]} (出现{len(n_values)}次)')
        else:
            print(f'  ⚠️ 多个不同的样本量: {unique_ns}')
            for n in unique_ns:
                count = n_values.count(n)
                print(f'     N={n} 出现 {count} 次')
    else:
        print(f'  ℹ️ 未检测到 N= 格式的样本量')
    
    # 4. 检查P值范围
    print(f'\n─── P值合理性检查 ───')
    p_pattern = re.findall(r'[pP]\s*[=<>＝＜＞]\s*(0\.\d+)', all_text + all_table_text)
    if p_pattern:
        p_values = [float(p) for p in p_pattern]
        invalid_p = [p for p in p_values if p > 1 or p < 0]
        if invalid_p:
            print(f'  ❌ 发现无效P值: {invalid_p}')
        else:
            sig = [p for p in p_values if p < 0.05]
            nonsig = [p for p in p_values if p >= 0.05]
            print(f'  ✅ 所有P值在有效范围 (共{len(p_values)}个)')
            print(f'     显著(P<0.05): {len(sig)}个')
            print(f'     不显著(P≥0.05): {len(nonsig)}个')
    
    # 5. 如果提供了数据文件，做交叉验证
    if data_path and os.path.exists(data_path):
        print(f'\n─── 数据文件交叉验证 ───')
        df = pd.read_excel(data_path)
        print(f'  数据文件: {data_path}')
        print(f'  Shape: {df.shape}')
        
        # 检查数据文件的样本量是否在报告中出现
        n_data = len(df)
        if str(n_data) in all_text or str(n_data) in all_table_text:
            print(f'  ✅ 数据样本量 {n_data} 在报告中出现')
        else:
            print(f'  ⚠️ 数据样本量 {n_data} 未在报告中找到！')
    
    # 6. 检查表编号连续性
    print(f'\n─── 表编号连续性检查 ───')
    table_nums = re.findall(r'表\s*(\d+)', all_text)
    if table_nums:
        nums = sorted(set(int(n) for n in table_nums))
        expected = list(range(nums[0], nums[-1] + 1))
        missing = set(expected) - set(nums)
        if missing:
            print(f'  ⚠️ 表编号跳号! 缺少: {sorted(missing)}')
        else:
            print(f'  ✅ 表编号连续: 表{nums[0]}-表{nums[-1]}')
    
    print(f'\n═══ 验证完成 ═══')

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python scripts/verify_report.py "报告.docx" ["数据.xlsx"]')
        sys.exit(1)
    
    docx_path = sys.argv[1]
    data_path = sys.argv[2] if len(sys.argv) > 2 else None
    verify_report(docx_path, data_path)
