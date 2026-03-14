# -*- coding: utf-8 -*-
"""
ACE 三线表 Word 生成器
========================
提供可复用的三线表函数库，支持:
  - 标准三线表（描述性统计、交叉表）
  - 方差分析三线表（组别×时间 + LSD字母 + ANOVA F值）
  - 回归结果三线表

所有函数返回 python-docx Table 对象，可自由组合到文档中。

用法（作为模块导入）:
    from three_line_table import ThreeLineTable, create_doc_with_table
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml


class ThreeLineTable:
    """三线表工具类"""

    # ── 字体常量 ──
    FONT_CN = '宋体'
    FONT_EN = 'Times New Roman'
    FONT_TITLE_CN = '黑体'
    SIZE_BODY = 9       # 小五
    SIZE_NOTE = 8
    SIZE_TITLE = 10.5   # 小四

    @staticmethod
    def clear_borders(table):
        """清除表格所有边框"""
        tblPr = table._tbl.tblPr or table._tbl._add_tblPr()
        borders = parse_xml(
            '<w:tblBorders %s>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>' % nsdecls('w'))
        for existing in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(existing)
        tblPr.append(borders)

    @staticmethod
    def set_row_border(row, position, sz=12, val="single", color="000000"):
        """设置某行某位置的边框 (top/bottom)"""
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.tcPr or tc._add_tcPr()
            borders = tcPr.find(qn('w:tcBorders'))
            if borders is None:
                borders = parse_xml('<w:tcBorders %s/>' % nsdecls('w'))
                tcPr.append(borders)
            el = parse_xml(
                f'<w:{position} {nsdecls("w")} w:val="{val}" '
                f'w:sz="{sz}" w:space="0" w:color="{color}"/>')
            existing = borders.find(qn(f'w:{position}'))
            if existing is not None:
                borders.remove(existing)
            borders.append(el)

    @staticmethod
    def set_row_border_dashed(row, position, sz=4, color="000000"):
        """虚线边框"""
        ThreeLineTable.set_row_border(row, position, sz=sz,
                                       val="dashed", color=color)

    @staticmethod
    def set_cell(cell, text, font_cn=None, font_en=None, size=None,
                 bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        """设置单元格内容和字体"""
        if font_cn is None:
            font_cn = ThreeLineTable.FONT_CN
        if font_en is None:
            font_en = ThreeLineTable.FONT_EN
        if size is None:
            size = ThreeLineTable.SIZE_BODY

        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.15

        run = p.add_run(str(text))
        run.font.name = font_en
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 设置中文字体
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_cn)

    @staticmethod
    def merge_vertical(table, col, start_row, end_row):
        """纵向合并单元格"""
        if start_row < end_row:
            table.cell(start_row, col).merge(table.cell(end_row, col))

    @staticmethod
    def apply_three_lines(table, header_end_row=0, footer_start_row=None):
        """应用三线表边框: 顶粗线 + 表头底细线 + 底粗线"""
        ThreeLineTable.clear_borders(table)
        # 顶线（粗）
        ThreeLineTable.set_row_border(table.rows[0], 'top', sz=12)
        # 表头底线（细）
        ThreeLineTable.set_row_border(table.rows[header_end_row], 'bottom', sz=6)
        # 底线（粗）
        last_row = len(table.rows) - 1
        ThreeLineTable.set_row_border(table.rows[last_row], 'bottom', sz=12)

    @staticmethod
    def add_note(doc, text, font_cn=None, size=None):
        """在表格后添加注释段落"""
        if font_cn is None:
            font_cn = ThreeLineTable.FONT_CN
        if size is None:
            size = ThreeLineTable.SIZE_NOTE

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = ThreeLineTable.FONT_EN
        run.font.size = Pt(size)
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_cn)
        return p

    @staticmethod
    def add_table_title(doc, title_cn, title_en=None):
        """添加表标题（中英文）"""
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(2)
        tr = tp.add_run(title_cn)
        tr.font.name = ThreeLineTable.FONT_EN
        tr.font.size = Pt(ThreeLineTable.SIZE_TITLE)
        tr.font.bold = True
        rPr = tr._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
            tr._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), ThreeLineTable.FONT_CN)

        if title_en:
            te = doc.add_paragraph()
            te.alignment = WD_ALIGN_PARAGRAPH.CENTER
            te.paragraph_format.space_before = Pt(0)
            te.paragraph_format.space_after = Pt(4)
            ter = te.add_run(title_en)
            ter.font.name = ThreeLineTable.FONT_EN
            ter.font.size = Pt(9)

    @classmethod
    def build_simple(cls, doc, headers, rows, col_widths=None):
        """
        构建简单三线表（描述性统计等）

        Args:
            doc: Document 对象
            headers: list of str，表头
            rows: list of list，数据行
            col_widths: list of Cm，列宽（可选）
        Returns:
            table 对象
        """
        n_rows = 1 + len(rows)
        n_cols = len(headers)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for j, h in enumerate(headers):
            cls.set_cell(table.rows[0].cells[j], h, bold=True)

        # 数据
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cls.set_cell(table.rows[i + 1].cells[j], val)

        # 三线
        cls.apply_three_lines(table, header_end_row=0)

        # 列宽
        if col_widths:
            for row in table.rows:
                for j, w in enumerate(col_widths):
                    if j < n_cols:
                        row.cells[j].width = w

        return table

    @classmethod
    def build_regression(cls, doc, models, dep_var_name="因变量"):
        """
        构建回归结果三线表

        Args:
            doc: Document
            models: list of dict, 每个模型包含:
                {
                    'name': '模型1',
                    'variables': [{'name': 'X1', 'B': 0.23, 'SE': 0.05,
                                   'beta': 0.31, 't': 4.60, 'p': 0.001}],
                    'R2': 0.45, 'adj_R2': 0.43, 'F': 12.3, 'F_p': 0.000,
                    'delta_R2': 0.12 (可选)
                }
        Returns:
            table 对象
        """
        headers = ['变量', 'B', 'SE', 'β', 't', 'p']
        all_rows = []

        for model in models:
            # 模型名行
            all_rows.append(('model_header', model['name']))
            for var in model['variables']:
                all_rows.append(('data', var))
            # R² 行
            r2_text = f"R²={model['R2']:.3f}"
            if 'adj_R2' in model:
                r2_text += f", adj R²={model['adj_R2']:.3f}"
            if 'delta_R2' in model:
                r2_text += f", ΔR²={model['delta_R2']:.3f}"
            if 'F' in model:
                r2_text += f", F={model['F']:.2f}"
            all_rows.append(('stats', r2_text))

        n_rows = 1 + len(all_rows)
        table = doc.add_table(rows=n_rows, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for j, h in enumerate(headers):
            cls.set_cell(table.rows[0].cells[j], h, bold=True)

        # 数据
        for i, (row_type, data) in enumerate(all_rows):
            row = table.rows[i + 1]
            if row_type == 'model_header':
                cls.set_cell(row.cells[0], data, bold=True)
                for j in range(1, len(headers)):
                    cls.set_cell(row.cells[j], '')
            elif row_type == 'data':
                cls.set_cell(row.cells[0], data['name'])
                cls.set_cell(row.cells[1], f"{data['B']:.3f}")
                cls.set_cell(row.cells[2], f"{data['SE']:.3f}")
                cls.set_cell(row.cells[3], f"{data['beta']:.3f}")
                cls.set_cell(row.cells[4], f"{data['t']:.3f}")
                p_str = f"{data['p']:.3f}" if data['p'] >= 0.001 else "<0.001"
                if data['p'] < 0.01:
                    p_str += "**"
                elif data['p'] < 0.05:
                    p_str += "*"
                cls.set_cell(row.cells[5], p_str)
            elif row_type == 'stats':
                cls.set_cell(row.cells[0], data,
                             align=WD_ALIGN_PARAGRAPH.LEFT)
                for j in range(1, len(headers)):
                    cls.set_cell(row.cells[j], '')

        cls.apply_three_lines(table, header_end_row=0)
        return table


def create_doc_landscape():
    """创建横向 A4 文档"""
    doc = Document()
    for section in doc.sections:
        section.orientation = 1
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
    return doc


def create_doc_portrait():
    """创建纵向 A4 文档"""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
    return doc


# ── CLI 演示 ──
if __name__ == "__main__":
    print("三线表工具库已加载。")
    print("用法示例:")
    print("  from three_line_table import ThreeLineTable, create_doc_portrait")
    print("  doc = create_doc_portrait()")
    print("  ThreeLineTable.add_table_title(doc, '表1 描述性统计')")
    print("  ThreeLineTable.build_simple(doc, ['变量','M','SD'], [['X1','3.2','0.8']])")
    print("  doc.save('output.docx')")
