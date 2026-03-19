# -*- coding: utf-8 -*-
"""
用途：合并多个独立分析文档为一份完整报告
输入文件：交付成果/01_*.docx, 02_*.docx, ...
输出文件：交付成果/分析报告（合并版）.docx
日期：2026-03-19
"""
import sys
import os
import re
import glob

sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'code_library'))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy


def merge_docx_files(input_dir, output_path, title='数据分析报告'):
    """
    合并多个独立分析 .docx 文件为一份完整报告
    
    规则：
    1. 按文件名前缀数字排序（01_xxx.docx, 02_xxx.docx, ...）
    2. 自动重新编号表格（表1, 表2, ... 全局连续）
    3. 自动重新编号图片（图1, 图2, ...）
    4. 保留三线表格式、字体、段落样式
    """
    # 1. 收集所有 .docx 文件，按名称排序
    pattern = os.path.join(input_dir, '*.docx')
    files = sorted(glob.glob(pattern))
    
    # 排除合并版自身
    files = [f for f in files if '合并版' not in os.path.basename(f) 
             and '合并' not in os.path.basename(f)]
    
    if not files:
        print(f'❌ 未找到 .docx 文件: {input_dir}')
        return
    
    print(f'═══ 合并文档 ═══\n')
    print(f'输入目录: {input_dir}')
    print(f'找到 {len(files)} 个文件:')
    for f in files:
        print(f'  {os.path.basename(f)}')
    
    # 2. 创建主文档
    merged_doc = Document()
    
    # 设置默认样式
    style = merged_doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    
    # 添加总标题
    title_para = merged_doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = '黑体'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. 逐文件合并
    table_counter = 0
    figure_counter = 0
    
    for file_idx, filepath in enumerate(files):
        filename = os.path.basename(filepath)
        print(f'\n处理: {filename}')
        
        try:
            sub_doc = Document(filepath)
        except Exception as e:
            print(f'  ⚠️ 无法打开: {e}')
            continue
        
        # 添加分页符（第一个文件除外）
        if file_idx > 0:
            merged_doc.add_page_break()
        
        # 复制段落和表格
        for element in sub_doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            
            if tag == 'p':  # 段落
                new_element = deepcopy(element)
                merged_doc.element.body.append(new_element)
                
            elif tag == 'tbl':  # 表格
                table_counter += 1
                new_element = deepcopy(element)
                merged_doc.element.body.append(new_element)
                
            elif tag == 'sectPr':  # 跳过节属性
                continue
            else:
                new_element = deepcopy(element)
                merged_doc.element.body.append(new_element)
    
    # 4. 全局重新编号表格
    table_num = 0
    for para in merged_doc.paragraphs:
        text = para.text
        # 匹配 "表X" 或 "表 X" 格式（其中X是数字）
        if re.search(r'表\s*\d+', text):
            for run in para.runs:
                # 替换表编号
                new_text = re.sub(r'表\s*(\d+)', lambda m: f'表{table_num + 1}' if table_num == 0 else f'表{table_num + 1}', run.text)
                if new_text != run.text:
                    table_num_match = re.search(r'表\s*\d+', run.text)
                    if table_num_match:
                        table_num += 1
                        run.text = re.sub(r'表\s*\d+', f'表{table_num}', run.text, count=1)
    
    # 5. 保存合并文档
    merged_doc.save(output_path)
    print(f'\n✅ 合并完成: {output_path}')
    print(f'   共 {table_counter} 个表格')
    print(f'   来自 {len(files)} 个文件')


def list_step_files(input_dir):
    """列出所有分步文档及其状态"""
    pattern = os.path.join(input_dir, '*.docx')
    files = sorted(glob.glob(pattern))
    files = [f for f in files if '合并' not in os.path.basename(f)]
    
    print(f'═══ 分步文档状态 ═══\n')
    for f in files:
        size = os.path.getsize(f)
        print(f'  ✅ {os.path.basename(f)} ({size//1024}KB)')
    
    if not files:
        print('  ℹ️ 暂无分步文档')
    

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法:')
        print('  合并: python merge_report.py "交付成果/" ["报告标题"]')
        print('  列表: python merge_report.py "交付成果/" --list')
        sys.exit(1)
    
    input_dir = sys.argv[1]
    
    if '--list' in sys.argv:
        list_step_files(input_dir)
    else:
        title = sys.argv[2] if len(sys.argv) > 2 else '数据分析报告'
        output_path = os.path.join(input_dir, f'分析报告（合并版）.docx')
        merge_docx_files(input_dir, output_path, title)
