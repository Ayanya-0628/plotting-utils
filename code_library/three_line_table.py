# -*- coding: utf-8 -*-
# ace 代码库: three_line_table.py
# 从 SKILL.md 提取的可复用代码模板
# 使用时复制对应函数/代码段，替换变量名即可

# ══════ python-docx 三线表核心代码 ══════

from docx.oxml.ns import nsdecls, qn

from docx.oxml import parse_xml

from docx.shared import Pt, Cm, RGBColor

def clear_table_borders(table):

    tblPr = table._tbl.tblPr or table._tbl._add_tblPr()

    borders = parse_xml(

        '<w:tblBorders %s>'

        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'

        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'

        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'

        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'

        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'

        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'

        '</w:tblBorders>' % nsdecls('w'))

    for existing in tblPr.findall(qn('w:tblBorders')):

        tblPr.remove(existing)

    tblPr.append(borders)

def set_row_border(row, position, sz=12, val="single", color="000000"):

    for cell in row.cells:

        tc = cell._tc

        tcPr = tc.tcPr or tc._add_tcPr()

        borders = tcPr.find(qn('w:tcBorders'))

        if borders is None:

            borders = parse_xml('<w:tcBorders %s/>' % nsdecls('w'))

            tcPr.append(borders)

        el = parse_xml(f'<w:{position} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')

        existing = borders.find(qn(f'w:{position}'))

        if existing is not None:

            borders.remove(existing)

        borders.append(el)

def set_cell_font(cell, text, font_cn='宋体', font_en='Times New Roman', size=9, bold=False):

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ''

    p = cell.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(str(text))

    run.font.name = font_en

    run.font.size = Pt(size)

    run.font.bold = bold

    run.font.color.rgb = RGBColor(0, 0, 0)

    rPr = run._element.find(qn('w:rPr'))

    if rPr is None:

        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))

        run._element.insert(0, rPr)

    rFonts = rPr.find(qn('w:rFonts'))

    if rFonts is None:

        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))

        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), font_cn)


# ══════ 图题格式 ══════

fig_caption = doc.add_paragraph()

fig_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = fig_caption.add_run(f'图{fig_num} {caption_text}')

run.font.name = 'Times New Roman'

run.font.size = Pt(8)

# 设置东亚字体为宋体（同 set_cell_font 方法）


