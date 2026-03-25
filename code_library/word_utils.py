# -*- coding: utf-8 -*-
"""
ace 代码库: word_utils.py
用途：Word 报告生成的基础函数，三线表、正文段落、注释行、图题等
日期：2026-03-19 (v2.0 新增)

★ 所有 Word 生成的项目都应导入此文件，不要每次从零手写 ★
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


# ══════ 三线表核心函数 ══════

def clear_table_borders(table):
    """清除表格所有边框"""
    tblPr = table._tbl.tblPr or table._tbl._add_tblPr()
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>' % nsdecls('w'))
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(borders)


def set_row_border(row, position, sz=12, val="single", color="000000"):
    """设置行边框（position: 'top'/'bottom'）"""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.tcPr or tc._add_tcPr()
        borders = tcPr.find(qn('w:tcBorders'))
        if borders is None:
            borders = parse_xml('<w:tcBorders %s/>' % nsdecls('w'))
            tcPr.append(borders)
        el = parse_xml(
            f'<w:{position} {nsdecls("w")} w:val="{val}" '
            f'w:sz="{sz}" w:space="0" w:color="{color}"/>')
        existing = borders.find(qn(f'w:{position}'))
        if existing is not None:
            borders.remove(existing)
        borders.append(el)


def set_cell_font(cell, text, font_cn='宋体', font_en='Times New Roman',
                  size=9, bold=False, align='center'):
    """设置单元格文字及字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == 'center'
                   else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Pt(0)

    run = p.add_run(str(text))
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 设置东亚字体
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)


def format_table_cells(table):
    """统一格式化表格所有单元格（行距、对齐、间距）"""
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.first_line_indent = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_three_line_table(doc, headers, data_rows, title=None, col_widths=None):
    """添加标准三线表

    Args:
        doc: Document 对象
        headers: 表头列表，如 ['变量', '均值', '标准差', 'P值']
        data_rows: 数据行列表，如 [['年龄', '35.2', '8.1', '<0.001'], ...]
        title: 表标题，如 '表1 人口学特征'
        col_widths: 列宽列表（Inches），None=自动

    Returns:
        table: python-docx Table 对象
    """
    # 表标题
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(3)
        run = title_para.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.bold = True
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '宋体')

    # 创建表格
    n_rows = 1 + len(data_rows)
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)

    # 清除边框 + 三线
    clear_table_borders(table)
    set_row_border(table.rows[0], 'top', sz=12)      # 顶粗线
    set_row_border(table.rows[0], 'bottom', sz=6)     # 栏目细线
    set_row_border(table.rows[-1], 'bottom', sz=12)   # 底粗线

    # 表头
    for j, header in enumerate(headers):
        set_cell_font(table.rows[0].cells[j], header, bold=True)

    # 数据行
    for i, row_data in enumerate(data_rows):
        for j, val in enumerate(row_data):
            set_cell_font(table.rows[i + 1].cells[j], val)

    # 列宽
    if col_widths:
        for j, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Inches(width)

    # 统一格式
    format_table_cells(table)

    return table


# ══════ 正文段落 ══════

def add_body_text(doc, text, indent=True):
    """添加正文段落（宋体小四，首行缩进2字符）

    Args:
        doc: Document 对象
        text: 段落文字
        indent: 是否首行缩进
    """
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10.5)  # 小四
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    return para


# ══════ 表注 ══════

def add_note(doc, text):
    """添加表注/数据来源说明（五号宋体）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0

    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10.5)  # 五号
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    return para


# ══════ 一级标题 ══════

def add_heading(doc, text, level=1):
    """添加标题（黑体）"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


# ══════ 图题 ══════

def add_figure_caption(doc, fig_num, caption_text):
    """添加图题（图X 描述文字，8pt宋体居中）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f'图{fig_num} {caption_text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    return para


# ══════ 插入图片 ══════

def add_figure(doc, image_path, width_inches=5.5):
    """插入图片（居中，默认5.5英寸宽）"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return para


# ══════ 创建标准报告文档 ══════

def create_report_doc():
    """创建标准格式的 Word 文档（已设置默认样式）"""
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    # 设东亚字体
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml('<w:rPr %s/>' % nsdecls('w'))
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s/>' % nsdecls('w'))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    return doc
