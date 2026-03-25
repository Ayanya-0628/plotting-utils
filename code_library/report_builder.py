# -*- coding: utf-8 -*-
"""
ace 代码库: report_builder.py
用途：全盘重算的脚本架构模板 + 数据一致性验证
日期：2026-03-19 (v2.0 新增)

★ 这是 ace v2.0 的核心文件 ★
★ 解决"迭代更新时新旧数据混用"的根本问题 ★
★ 参见 SKILL.md §18 全盘重算铁律 ★

使用方式：
    新建分析项目时，复制本文件到项目目录，
    修改 CONFIG、load_and_clean()、analyze()、generate_report() 即可。
"""
import pandas as pd
import numpy as np
from docx import Document


# ══════════════════════════════════════════════
# 第0层：配置（修改这里适配你的项目）
# ══════════════════════════════════════════════

CONFIG = {
    'input_file': '原始数据/xxx.xlsx',
    'output_docx': '交付成果/分析结果.docx',
    'output_xlsx': '交付成果/分析数据.xlsx',
    'project_name': 'XXX项目',
}


# ══════════════════════════════════════════════
# 第1层：数据加载与清洗
# ══════════════════════════════════════════════

def load_and_clean(config):
    """从原始数据加载，返回清洗后的 DataFrame
    
    ★ 这是唯一读取原始数据的地方 ★
    ★ 后续层不允许重新读取数据 ★
    """
    df = pd.read_excel(config['input_file'])
    
    # TODO: 反向计分
    # TODO: 缺失值处理
    # TODO: 编码校验
    # TODO: 维度得分计算
    
    print(f'✅ 数据加载完成: N={len(df)}, 变量数={df.shape[1]}')
    return df


# ══════════════════════════════════════════════
# 第2层：统计分析
# ══════════════════════════════════════════════

def analyze(df):
    """所有统计分析，返回 results 字典
    
    ★ 所有统计结果必须存入 results 字典 ★
    ★ 此函数不做任何 Word/Excel 输出 ★
    ★ 只接受 df 作为输入，不引用外部变量 ★
    """
    results = {}
    
    # --- 人口学特征 ---
    # results['demographic'] = {...}
    
    # --- 描述性统计 ---
    # results['descriptive'] = {...}
    
    # --- 信度检验 ---
    # results['reliability'] = {...}
    
    # --- 效度检验 ---
    # results['validity'] = {...}
    
    # --- 相关分析 ---
    # results['correlation'] = {...}
    
    # --- 差异分析 ---
    # results['difference'] = {...}
    
    # --- 回归分析 ---
    # results['regression'] = {...}
    
    # --- 中介效应 ---
    # results['mediation'] = {...}
    
    print(f'✅ 统计分析完成: {len(results)} 个模块')
    return results


# ══════════════════════════════════════════════
# 第3层：报告生成
# ══════════════════════════════════════════════

def generate_report(results, config):
    """根据 results 字典生成 Word 报告
    
    ★ 所有表格和文字都从 results 字典中取值 ★
    ★ 禁止在此层重新计算任何统计量 ★
    ★ 禁止 import pandas / 重新读取数据文件 ★
    ★ 每个 add_three_line_table() 后面必须紧跟 add_body_text() ★
    """
    # from word_utils import (create_report_doc, add_heading, 
    #     add_three_line_table, add_body_text, add_note)
    
    doc = Document()  # 或 create_report_doc()
    
    table_num = 0  # 表编号自增，防跳号
    
    # --- 示例：人口学特征 ---
    # table_num += 1
    # add_heading(doc, '一、人口学特征', level=1)
    # add_body_text(doc, '本研究共回收有效问卷...份，...')
    # add_three_line_table(doc, headers, rows, title=f'表{table_num} ...')
    # add_note(doc, '注：...')
    # add_body_text(doc, '由表{table_num}可知，...')  # 紧跟文字分析！
    
    doc.save(config['output_docx'])
    print(f'✅ Word报告已生成: {config["output_docx"]}')


# ══════════════════════════════════════════════
# 验证层：数据一致性检查
# ══════════════════════════════════════════════

def verify_report(docx_path, expected_checks):
    """验证 Word 报告中的关键数字是否与预期一致
    
    Args:
        docx_path: Word 文件路径
        expected_checks: [(标签, 期望值), ...] 列表
    
    示例:
        verify_report('交付成果/分析结果.docx', [
            ('样本量', '300'),
            ('男性人数', '156'),
            ('总量表alpha', '0.923'),
        ])
    """
    doc = Document(docx_path)
    
    # 提取所有文字（段落+表格）
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += '\n' + cell.text
    
    errors = []
    passed = 0
    for label, expected in expected_checks:
        if str(expected) not in all_text:
            errors.append(f'  ❌ {label}: 期望 [{expected}] 未在报告中找到')
        else:
            passed += 1
    
    if errors:
        print(f'数据一致性验证: {passed}/{len(expected_checks)} 通过')
        print('\n'.join(errors))
        raise ValueError(
            f'数据一致性验证失败！{len(errors)}处不一致，请检查！')
    else:
        print(f'✅ 全部 {len(expected_checks)} 项验证通过')


# ══════════════════════════════════════════════
# 主入口（唯一入口）
# ══════════════════════════════════════════════

if __name__ == '__main__':
    # ★ 从头到尾一次性运行，禁止分段执行 ★
    # ★ 任何修改后，必须从头重新运行整个脚本 ★
    
    df = load_and_clean(CONFIG)         # 第1层：加载
    results = analyze(df)               # 第2层：分析
    generate_report(results, CONFIG)    # 第3层：生成
    
    # 可选：数据一致性验证
    # verify_report(CONFIG['output_docx'], [
    #     ('样本量', str(len(df))),
    #     # 添加更多检查项...
    # ])
